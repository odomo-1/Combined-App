import streamlit as st
import os
import tempfile
import docx
from docx import Document
import re
from io import BytesIO
from spellchecker import SpellChecker
import pandas as pd
import spacy
from pdfminer.high_level import extract_text

# --- Constants ---
STANDARD_SECTIONS = [
    "Table of content",
    "Introduction",
    "Background",
    "Objective",
    "Methodology" or "Approach",
    "Project Team",
    "About Sahel",
    "Budget",
    "Work Plan",
]

# --- Helper Functions (from both apps) ---

def extract_text_from_pdf(pdf_file):
    return extract_text(pdf_file)

def extract_text_from_word(word_file):
    doc = Document(word_file)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def clean_text(text):
    return re.sub(r"[\x00-\x1F\x7F]", "", text)

def extract_sentences_with_keywords(text, keywords, assigned_sentences):
    sentences = re.split(r'(?<=[.!?])\s+', text)
    keyword_pattern = re.compile(r'(?i)\b(?:' + '|'.join(keywords) + r')\b')
    matches = []

    for sentence in sentences:
        if keyword_pattern.search(sentence) and sentence.strip() not in assigned_sentences:
            matches.append(sentence.strip())
            assigned_sentences.add(sentence.strip())

    return matches if matches else ["Not Found"]

def extract_named_entities(text, nlp, label, assigned_sentences):
    doc = nlp(text)
    matches = []

    for ent in doc.ents:
        if ent.label_ == label and ent.text not in assigned_sentences:
            matches.append(ent.text)
            assigned_sentences.add(ent.text)

    return matches if matches else ["Not Found"]

def categorize_rfp(text):
    grant_keywords = ["grant", "funding", "donation", "philanthropy", "financial aid"]
    investment_keywords = ["investment", "capital", "funding", "venture", "equity"]
    assessment_keywords = ["assessment", "evaluation", "review", "impact", "audit"]
    market_research_keywords = ["market research", "consumer research", "market analysis", "industry study", "market survey"]

    if any(re.search(r'\b' + keyword + r'\b', text, re.IGNORECASE) for keyword in grant_keywords):
        return "Grant"
    elif any(re.search(r'\b' + keyword + r'\b', text, re.IGNORECASE) for keyword in investment_keywords):
        return "Investment"
    elif any(re.search(r'\b' + keyword + r'\b', text, re.IGNORECASE) for keyword in assessment_keywords):
        return "Assessment"
    elif any(re.search(r'\b' + keyword + r'\b', text, re.IGNORECASE) for keyword in market_research_keywords):
        return "Market Research"
    else:
        return "Uncategorized"

def process_rfp(file, file_type):
    if file_type == "pdf":
        text = extract_text_from_pdf(file)
    elif file_type == "docx":
        text = extract_text_from_word(file)
    else:
        raise ValueError("Unsupported file type")

    text = clean_text(text)
    nlp = spacy.load("./en_core_web_sm")

    rfp_category = categorize_rfp(text)

    scope_keywords = ["Scope", "Description", "Objective", "Goals", "Deliverables", "Statement of Work"]
    methodology_keywords = ["Methodology", "Approach", "Strategy", "Plan", "Implementation", "Execution", "Framework", "Process", "Techniques", "Procedures"]
    eligibility_keywords = ["Eligibility", "Eligible", "Applicants", "Who can apply", "Requirements", "Qualifications", "Criteria", "Conditions", "Target Audience"]
    budget_keywords = ["Budget", "Funding", "Cost", "Financial", "Expenses", "Price", "Pricing", "Allocation", "Payment Terms"]
    deadline_keywords = ["Deadline", "Submission", "Due Date", "Timeline", "Schedule", "Important Dates"]
    selection_process_keywords = ["Selection", "Evaluation", "Criteria", "Process", "Weighting", "Judging", "Metrics", "Assessment", "Decision"]

    assigned_sentences = set()

    extracted_info = {
        "Section": [
            "Scope of Work", "Methodology", "Eligibility",
            "Budget", "Deadlines", "Selection Process"
        ],
        "Details": [
            "\n".join(extract_sentences_with_keywords(text, scope_keywords, assigned_sentences)),
            "\n".join(extract_sentences_with_keywords(text, methodology_keywords, assigned_sentences)),
            "\n".join(extract_sentences_with_keywords(text, eligibility_keywords, assigned_sentences)),
            "\n".join(extract_sentences_with_keywords(text, budget_keywords, assigned_sentences)),
            "\n".join(extract_named_entities(text, nlp, "DATE", assigned_sentences)),
            "\n".join(extract_sentences_with_keywords(text, selection_process_keywords, assigned_sentences)),
        ]
    }

    df = pd.DataFrame(extracted_info)
    return rfp_category, df

def save_to_word(rfp_category, df):
    doc = Document()
    doc.add_heading("RFP Extracted Information", level=1)
    doc.add_heading("RFP Category", level=2)
    doc.add_paragraph(rfp_category, style="BodyText")

    for index, row in df.iterrows():
        doc.add_heading(row["Section"], level=2)
        doc.add_paragraph(row["Details"], style="BodyText")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def extract_text(file):
    text = ""
    if file.name.endswith('.docx'):
        temp_path = os.path.join(tempfile.gettempdir(), file.name)
        with open(temp_path, 'wb') as f:
            f.write(file.read())
        doc = Document(temp_path)
        for para in doc.paragraphs:
            text += para.text + '\n'
    return text

# --- Budget Check Logic ---
def budget_check(doc):
    standard_rates = {
        "Project Director": 1400,
        "Project Manager": 1200,
        "Consultant": 850,
        "Analyst": 700
    }

    budget_section = ""
    for para in doc.paragraphs:
        if "budget" in para.text.lower():
            budget_section += para.text + "\n"

    budget_text = budget_section.lower()

    role_keywords = ["project director", "project manager", "consultant", "analyst"]
    role_rates = {role: standard_rates[role] for role in role_keywords}

    rate_mismatches = []

    for role, standard_rate in role_rates.items():
        if role in budget_text:
            match = re.search(rf"({role}).*?(\d+)", budget_text)
            if match:
                role_amount = int(match.group(2))
                if role_amount != standard_rate:
                    rate_mismatches.append(f"{role.capitalize()} rate of ${role_amount} does not match standard of ${standard_rate}.")

    return rate_mismatches

def evaluate_proposal(text, required_sections, doc):
    lower_text = text.lower()

    section_results = {}
    for sec in required_sections:
        found = any(sec.lower() in para.text.lower() for para in doc.paragraphs)
        section_results[sec] = found

    section_score = sum(section_results.values())
    section_percentage = (section_score / len(required_sections)) * 100

    formatting_results = formatting_check(doc)
    budget_issues = budget_check(doc)

    total_score = 0
    max_score = 4

    total_score += section_percentage * 0.50

    spell_score = 0
    if len(formatting_results['spelling_issues']) == 0:
        spell_score = 100
    else:
        spell_score = max(0, 100 - len(formatting_results['spelling_issues']) * 10)
    total_score += spell_score * 0.25

    font_style_score = 100 if formatting_results['font_ok'] else 0
    font_size_score = 100 if formatting_results['font_size_ok'] else 0
    total_score += (font_style_score + font_size_score) * 0.25

    total_score = round(total_score)

    # Recommendations
    missing_sections = [sec for sec, present in section_results.items() if not present]
    recommendations = []
    if missing_sections:
        recommendations.append(f"Kindly include the following missing sections: {', '.join(missing_sections)}")
    if formatting_results['spelling_issues']:
        recommendations.append("Spelling issues found in the document.")
    if not formatting_results['font_ok']:
        recommendations.append("Document should use font 'Tenorite' throughout.")
    if not formatting_results['font_size_ok']:
        recommendations.append("Body text should use font size 11.")
    if budget_issues:
        recommendations.extend(budget_issues)

    return {
        'sections': section_results,
        'score': total_score,
        'recommendations': recommendations,
        'formatting': formatting_results
    }

# --- Streamlit App Interface ---
st.title("Proposal Toolkit")

app_mode = st.radio("Select Tool", ["Proposal Evaluator", "RFP Key Info Extractor"])

if app_mode == "Proposal Evaluator":
    uploaded_proposal = st.file_uploader("Upload Proposal (.docx only)", type=["docx"])
    evaluation = None

    if uploaded_proposal and st.button("Evaluate Proposal"):
        st.success("Proposal uploaded successfully.")
        prop_text = extract_text(uploaded_proposal)
        doc = Document(uploaded_proposal)
        with st.spinner("Evaluating proposal..."):
            evaluation = evaluate_proposal(prop_text, STANDARD_SECTIONS, doc)

    if evaluation:
        st.subheader("Evaluation Results")
        st.write("### Section Check")
        for section, found in evaluation['sections'].items():
            st.write(f"- **{section}**: {'✅' if found else '❌'}")

        st.write("### Formatting & Presentation")
        if evaluation['formatting']['spelling_issues']:
            st.warning("Spelling Issues Detected:")
            st.write(", ".join(evaluation['formatting']['spelling_issues']))
        else:
            st.success("No major spelling issues detected.")

        if evaluation['formatting']['font_ok'] and evaluation['formatting']['font_size_ok']:
            st.success("Font style and size meet organizational standards (Tenorite, size 11).")
        else:
            st.warning("Font style does not match standard (Tenorite) or font size is not 11 in body text.")

        st.write(f"### Overall Score: **{evaluation['score']}%**")

        st.write("### Recommendations")
        if evaluation['recommendations']:
            for rec in evaluation['recommendations']:
                st.warning(rec)
        else:
            st.success("Your proposal aligns well with the standards!")

        word_buffer = create_word_report(evaluation)
        st.download_button(
            label="Download Evaluation Report (.docx)",
            data=word_buffer,
            file_name="proposal_evaluation.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

elif app_mode == "RFP Key Info Extractor":
    uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx"])
    if uploaded_file is not None:
        file_type = uploaded_file.name.split('.')[-1].lower()
        if file_type not in ["pdf", "docx"]:
            st.error("Unsupported file type. Please upload a PDF or Word document.")
        else:
            rfp_category, df = process_rfp(uploaded_file, file_type)
            st.write("### RFP Category")
            st.success(rfp_category)
            st.write("### Extracted Information")
            st.dataframe(df)

            buffer = save_to_word(rfp_category, df)
            st.download_button(
                "Download Extracted Info as Word",
                buffer,
                file_name="rfp_extracted_info.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
